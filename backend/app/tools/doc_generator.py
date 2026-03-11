"""Word document generation builtin tool.

Input params:
{
  "title": "Q1季度营销方案",
  "sections": [
    {"type": "heading", "level": 1, "text": "一、项目背景"},
    {"type": "paragraph", "text": "本季度计划..."},
    {"type": "paragraph", "text": "重点关注", "bold": true},
    {"type": "list", "items": ["要点一", "要点二"]},
    {"type": "table", "headers": ["指标", "Q4", "Q1目标"], "rows": [["GMV", "500万", "800万"]]}
  ]
}

Output: {"file_id": "xxx.docx", "download_url": "/api/files/xxx.docx", "filename": "Q1季度营销方案.docx"}
"""
from __future__ import annotations

import os
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

_UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "./uploads")
_GENERATED_DIR = Path(_UPLOAD_DIR) / "generated"


def execute(params: dict) -> dict:
    """Generate a DOCX file and return file_id + download_url."""
    _GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    title_text = params.get("title", "文档")
    sections = params.get("sections", [])

    doc = Document()

    # Cover title
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title_text)
    run.bold = True
    run.font.size = Pt(24)
    doc.add_paragraph()  # spacer

    for sec in sections:
        sec_type = sec.get("type", "paragraph")

        if sec_type == "heading":
            level = sec.get("level", 1)
            level = max(1, min(3, level))
            doc.add_heading(sec.get("text", ""), level=level)

        elif sec_type == "paragraph":
            p = doc.add_paragraph()
            run = p.add_run(sec.get("text", ""))
            if sec.get("bold"):
                run.bold = True
            if sec.get("italic"):
                run.italic = True

        elif sec_type == "list":
            for item in sec.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(item))

        elif sec_type == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            if headers:
                col_count = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=col_count)
                table.style = "Table Grid"

                # Header row with background color
                hdr_row = table.rows[0]
                for i, h in enumerate(headers):
                    cell = hdr_row.cells[i]
                    cell.text = h
                    # Set header background color (#4472C4)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "4472C4")
                    tcPr.append(shd)
                    # White bold font
                    para = cell.paragraphs[0]
                    run = para.runs[0] if para.runs else para.add_run(h)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    row_obj = table.rows[r_idx + 1]
                    for c_idx, val in enumerate(row_data[:col_count]):
                        row_obj.cells[c_idx].text = str(val)

                # Set auto-fit
                table.autofit = True
                doc.add_paragraph()  # spacer after table

    file_id = f"{uuid.uuid4().hex}.docx"
    file_path = _GENERATED_DIR / file_id
    doc.save(str(file_path))

    return {
        "file_id": file_id,
        "download_url": f"/api/files/{file_id}",
        "filename": f"{title_text}.docx",
        "message": f"文档已生成：{title_text}，共{len(sections)}个章节",
    }
